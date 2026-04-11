
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime, os

doc = Document()

# PAGE SETUP (A4)
section = doc.sections[0]
section.page_width  = Inches(8.27)
section.page_height = Inches(11.69)
section.left_margin = section.right_margin = Inches(0.8)
section.top_margin  = section.bottom_margin = Inches(0.8)

IMG_DIR = r"d:\NEWNUSTECH\manual_screenshots"
IMG_W   = Inches(6.0)

# ── helper funcs ─────────────────────────────────────────────────────────────

def h1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text.upper())
    r.font.size, r.font.bold = Pt(16), True
    r.font.color.rgb = RGBColor(0x07,0x11,0x52)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(8)
    return p

def h_group(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text.upper())
    r.font.size, r.font.bold = Pt(22), True
    r.font.color.rgb = RGBColor(0x0d,0x6e,0xfd)
    p.paragraph_format.space_before = Pt(30)
    p.paragraph_format.space_after  = Pt(14)
    return p

def h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size, r.font.bold = Pt(14), True
    r.font.color.rgb = RGBColor(0x0d,0x6e,0xfd)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    return p

def h3(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size, r.font.bold = Pt(11), True
    r.font.color.rgb = RGBColor(0x19,0x87,0x54)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(4)

def body(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor(0x2c,0x3e,0x50)
    p.paragraph_format.space_after = Pt(5)

def step(doc, n, text):
    p = doc.add_paragraph()
    rn = p.add_run(f"{n}. ")
    rn.font.bold, rn.font.color.rgb = True, RGBColor(0x0d,0x6e,0xfd)
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(3)

def bullet(doc, text, prefix=None):
    p = doc.add_paragraph()
    p.add_run("• ").font.bold = True
    if prefix:
        p.add_run(prefix + ": ").font.bold = True
    p.add_run(text)
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(2)

def filter_detail(doc, name, desc):
    p = doc.add_paragraph()
    r = p.add_run(f"   - {name}: ")
    r.font.bold = True
    p.add_run(desc)
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(2)

def note(doc, text):
    p = doc.add_paragraph()
    r = p.add_run("📌 Catatan: ")
    r.font.bold, r.font.color.rgb = True, RGBColor(0xe6,0x7e,0x22)
    p.add_run(text).font.size = Pt(10)
    p.paragraph_format.left_indent = Inches(0.2)

def divider(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'12'); bot.set(qn('w:color'),'0d6efd')
    pBdr.append(bot); pPr.append(pBdr)

def add_img(doc, filename, caption=None):
    path = os.path.join(IMG_DIR, filename)
    if not os.path.exists(path): return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(path, width=IMG_W)
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run(f"Gambar: {caption}")
        cr.font.italic, cr.font.size = True, Pt(9)
        cr.font.color.rgb = RGBColor(0x7f,0x83,0x91)

# ═══════════════════════════════════════════════════════════════════════════════
#  COVER & INTRODUCTION
# ═══════════════════════════════════════════════════════════════════════════════
for _ in range(3): doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("NUSTECH")
r.font.size, r.font.bold = Pt(44), True
r.font.color.rgb = RGBColor(0x07,0x11,0x52)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("PANDUAN LENGKAP PENGGUNAAN WEBSITE")
r.font.size, r.font.bold = Pt(20), True
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("MODUL KHUSUS ROLE: USER")
r.font.size, r.font.bold = Pt(14), True
r.font.color.rgb = RGBColor(0x19,0x87,0x54)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(f"Diperbarui Pada: {datetime.datetime.now().strftime('%d %B %Y')}")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("CV. NUSTECH — Lombok, NTB")
doc.add_page_break()

# 1. LANDING PAGE
# ─────────────────────────────────────────────────────────────────────────────
h1(doc, "MODUL 1: AKSES & MONITORING UTAMA")
divider(doc)

h2(doc, "1.1 Landing Page")
body(doc, "Landing page adalah pintu gerbang awal sistem. Ini memberikan gambaran umum layanan sebelum pengguna masuk ke area terproteksi.")
add_img(doc, "00_landing.png", "Halaman Utama (Landing Page)")
h3(doc, "Fitur Utama:")
bullet(doc, "Menampilkan jumlah total site, tiket, dan pencapaian operasional.", "Statistik Operasional")
bullet(doc, "Akses cepat ke portal Monitoring Network System (MNS).", "Quick Access")
step(doc, 1, "Buka peramban (browser) dan akses alamat website NUSTECH.")
step(doc, 2, "Klik tombol 'MASUK' atau 'LOGIN' untuk menuju halaman autentikasi.")

h2(doc, "1.2 Halaman Login")
body(doc, "Sistem keamanan NUSTECH menggunakan autentikasi email dan password terenkripsi.")
add_img(doc, "00_login_page.png", "Form Login Pengguna")
h3(doc, "Fitur Keamanan:")
bullet(doc, "Klik ikon mata untuk memastikan password yang diketik sudah benar.", "Show Password")
bullet(doc, "Centang opsi ini agar sesi tetap aktif dalam jangka waktu tertentu.", "Remember Me")
step(doc, 1, "Masukkan Email yang telah didaftarkan oleh admin.")
step(doc, 2, "Masukkan Password dengan benar (perhatikan huruf besar/kecil).")
step(doc, 3, "Klik 'MASUK' untuk mengakses Dashboard.")

h2(doc, "1.3 My Dashboard (Pusat Kendali)")
body(doc, "Halaman pertama setelah login yang memberikan pandangan 360 derajat terhadap kondisi operasional.")
add_img(doc, "00_dashboard.png", "Dashboard Monitoring Utama")
h3(doc, "Fitur Utama Dashboard:")
bullet(doc, "Kotak informasi di bagian atas (Tiket Hari Ini, Tiket Terbuka, Progres PM).", "Visual Stats")
bullet(doc, "Komunikasi pesan instan antar user sistem untuk koordinasi cepat.", "Live Chat Box")
bullet(doc, "Daftar teknisi yang sedang bertugas pada shift Pagi, Siang, atau Malam.", "Jadwal Shift Hari Ini")
bullet(doc, "Tabel yang menampilkan daftar kendala terbaru yang membutuhkan perhatian.", "Open Ticket Table Summary")

h3(doc, "Tombol Statistik (Stats Badges):")
add_img(doc, "stats_badges.png", "Tombol Statistik Dashboard")
filter_detail(doc, "Total Open", "Menampilkan jumlah seluruh tiket yang masih berstatus terbuka (belum selesai).")
filter_detail(doc, "Open Hari Ini", "Menampilkan jumlah tiket baru yang masuk pada tanggal hari ini.")
filter_detail(doc, "BMN / SL", "Memisahkan jumlah tiket berdasarkan kategori proyek bantuan pemerintah atau swasta.")
note(doc, "Anda dapat mengklik angka pada badge statistik ini untuk menyaring tabel secara otomatis.")

h2(doc, "1.4 Menu Navigasi (Daftar Halaman Operasional)")
body(doc, "Sistem memiliki menu 'Mega Menu' berupa modal popup untuk berpindah antar halaman secara cepat. Menu ini terbagi menjadi 4 pilar utama:")
add_img(doc, "halaman_operasional.png", "Modal Daftar Halaman Operasional (Mega Menu)")
h3(doc, "Rincian Kategori Menu:")
bullet(doc, "Berisi Data Site, Manajemen Password, Laporan CM (Corrective), Laporan PM (Preventive), dan Summary PM.", "1. Data Site")
bullet(doc, "Berisi Open Tiket, Close Tiket, Detail Tiket, dan Summary Tiket untuk analisis performa.", "2. Tiket")
bullet(doc, "Berisi Pergantian Perangkat, Log Pergantian (Riwayat), Spare Tracker, dan Summary Stok Perangkat.", "3. Log Perangkat")
bullet(doc, "Berisi tugas pribadi (Todo List), Jadwal Piket, dan audit koneksi (Log Remote).", "4. Project Info")
doc.add_page_break()

# 2. GROUP: DATA SITE
# ─────────────────────────────────────────────────────────────────────────────
h_group(doc, "MODUL 2: INVENTARIS & DATABASE SITE")
divider(doc)

h2(doc, "2.1 All Sites (Database Master Site)")
body(doc, "Halaman ini memuat seluruh data teknis lokasi site yang dikelola.")
add_img(doc, "01_datasite.png", "Halaman Data Site Terintegrasi")
h3(doc, "Fungsi Filter & Fitur:")
filter_detail(doc, "Pencarian (Search)", "Mencari site berdasarkan Nama Site, Site ID, atau Alamat secara fleksibel.")
filter_detail(doc, "Provinsi/Kabupaten", "Menyaring tampilan data hanya untuk wilayah tertentu (misal: Provinsi NTB, Kab. Lombok Barat).")
filter_detail(doc, "Tipe Site", "Membedakan site berdasarkan kategori proyek (misal: BMN atau SL).")
bullet(doc, "Kolom seperti IP Router dan IP AP terkunci di kiri agar tidak hilang saat tabel digeser horizontal.", "Sticky Columns")
bullet(doc, "Terdapat tombol Download untuk mengunduh seluruh data dalam format Excel.", "Export Data")

h2(doc, "2.2 Manajemen Password")
body(doc, "Tempat penyimpanan kredensial akses perangkat (mikrotik/AP) untuk setiap site.")
add_img(doc, "02_manajemen_pw.png", "Halaman Manajemen Password")
h3(doc, "Fungsi Filter:")
filter_detail(doc, "Filter Wilayah", "Sama dengan Data Site, memungkinkan pencarian password berdasarkan lokasi geografis.")
filter_detail(doc, "Search Box", "Mencari Nama Site secara spesifik untuk melihat password login-nya.")

h2(doc, "2.3 Corrective Maintenance (Laporan CM)")
body(doc, "Berfungsi untuk mendata perbaikan di lapangan yang bersifat insidental (gangguan).")
add_img(doc, "03_laporan_cm.png", "Tabel Laporan Corrective Maintenance")
h3(doc, "Fungsi Filter Spesifik:")
filter_detail(doc, "Kategori Laporan", "Dropdown untuk memfilter jenis CM (misal: Reset Perangkat, Ganti FO, dsb).")
filter_detail(doc, "Rentang Tanggal", "Mengatur periode tampilan laporan (dari tanggal X sampai tanggal Y).")
bullet(doc, "Anda dapat melihat bukti transfer atau foto on-site teknisi melalui link pada tabel.", "Verifikasi Visual")

h2(doc, "2.4 Summary PM & PM Liberta")
body(doc, "Halaman rekapitulasi kegiatan Preventive Maintenance (Perawatan Rutin).")
add_img(doc, "15_summarypm.png", "Chart & Statistik Summary PM")
h3(doc, "Analisis Data:")
bullet(doc, "Grafik batang menunjukkan progres penyelesaian PM per bulan.", "Visual Chart")
bullet(doc, "Daftar detail site mana saja yang sudah atau belum dilakukan perawatan rutin.", "PM Inventory")
doc.add_page_break()

# 3. GROUP: TIKET
# ─────────────────────────────────────────────────────────────────────────────
h_group(doc, "MODUL 3: SISTEM TIKET GANGGUAN")
divider(doc)

h2(doc, "3.1 Open & Close Tiket")
body(doc, "Inti dari sistem monitoring untuk menangani laporan downtime atau kendala teknis.")
add_img(doc, "06_open_tiket.png", "Monitoring Tiket Aktif (Open)")

h3(doc, "Detail Filter Bar (Panduan Satu Persatu):")
add_img(doc, "filter_bar.png", "Komponen Filter & Pencarian Standar")
filter_detail(doc, "Entries (Show X Data)", "Kotak angka di paling kiri (misal: 10, 50, 100) untuk menentukan berapa banyak baris data yang ditampilkan dalam satu halaman tabel.")
filter_detail(doc, "Dropdown Kategori", "Memilih jenis layanan spesifik (BMN / SL) agar tabel hanya menampilkan tiket dari kategori tersebut.")
filter_detail(doc, "Date Pickers (Rentang Tanggal)", "Dua kolom input kalender untuk melihat data dalam kurun waktu tertentu (Start Date s/d End Date).")
filter_detail(doc, "Tombol Filter (Biru)", "Tombol utama untuk menjalankan penyaringan data setelah Anda menentukan kategori atau tanggal.")
filter_detail(doc, "Tombol Refresh (Ikon Panah Putar)", "Berguna untuk mereset seluruh filter ke kondisi awal dan memperbarui data terbaru dari server.")
filter_detail(doc, "Search Input (Pencarian)", "Kolom teks di paling kanan untuk mencari Nama Site atau ID Site secara langsung (tanpa harus klik tombol filter).")

bullet(doc, "Klik 'Site ID' untuk membuka modal rincian site beserta lokasi Maps yang presisi.", "Quick Detail View")
add_img(doc, "08_detail_tiket.png", "Modal Detail Tiket & Map")
h3(doc, "Summary Tiket:")
add_img(doc, "06_summary_tiket.png", "Analisis Tren Tiket")
body(doc, "Melihat tren gangguan terbanyak, durasi penyelesaian rata-rata, dan statistik per kabupaten.")
doc.add_page_break()

# 4. GROUP: LOG PERANGKAT
# ─────────────────────────────────────────────────────────────────────────────
h_group(doc, "MODUL 4: LOG PERANGKAT & LOGISTIK")
divider(doc)

h2(doc, "4.1 Pergantian & Log Perangkat")
body(doc, "Mencatat sejarah perangkat yang terpasang di setiap site.")
add_img(doc, "10_pergantian.png", "Monitoring Pergantian Perangkat")
h3(doc, "Fitur Pelacakan:")
filter_detail(doc, "Search Serial Number", "Mencari posisi perangkat berdasarkan SN lama atau SN baru.")
filter_detail(doc, "Status Perangkat", "Filter untuk melihat perangkat yang Baru, Rusak, atau Sedang Ditarik.")

h2(doc, "4.2 Spare Tracker")
body(doc, "Melacak distribusi stok sparepart (modem, router, pigtail, dll).")
add_img(doc, "12_spare_tracker.png", "Logistik Spare Tracker")
filter_detail(doc, "Nama Sparepart", "Mencari ketersediaan atau posisi suatu jenis alat.")
doc.add_page_break()

# 5. GROUP: PROJECT INFO
# ─────────────────────────────────────────────────────────────────────────────
h_group(doc, "MODUL 5: PROJECT MANAGEMENT")
divider(doc)

h2(doc, "5.1 Sparepart Needed (Permintaan Alat)")
body(doc, "Formulir digital untuk mengajukan kebutuhan alat bantu atau sparepart.")
add_img(doc, "08_sparepart.png", "Daftar Pengajuan Sparepart")
h3(doc, "Alur Kerja:")
filter_detail(doc, "Filter Status", "Melihat status pengajuan (Pending, Approved, atau Completed).")
bullet(doc, "Terdapat fitur untuk mencetak formulir pengajuan resmi yang bisa ditandatangani.", "Print Formulir")
bullet(doc, "Setiap pengajuan dapat dilampiri foto resi atau foto barang yang diminta.", "Lampiran Foto")

h2(doc, "5.2 My Todo List")
body(doc, "Buku catatan digital untuk tugas individu pengguna.")
add_img(doc, "14_todolist.png", "Personal Task Manager")
h3(doc, "Cara Menggunakan:")
bullet(doc, "Ketik judul project besar di input utama.", "Master Task")
bullet(doc, "Gunakan input kecil di dalam kartu untuk memecah tugas menjadi langkah-langkah kecil.", "Sub-Tasking")
bullet(doc, "Progress bar akan terisi otomatis seiring banyaknya sub-task yang Anda centang.", "Auto Progress")

h2(doc, "5.3 Jadwal Piket (Shift Control)")
body(doc, "Transparansi jadwal kerja seluruh tim nustech.")
add_img(doc, "12_jadwal_piket.png", "Kalender Jadwal Shift")
h3(doc, "Fitur Pendukung:")
filter_detail(doc, "Bulan/Tahun", "Gunakan dropdown untuk melihat jadwal di masa lalu atau bulan depan.")
bullet(doc, "Gunakan tombol CAPTURE untuk menyimpan tampilan jadwal menjadi gambar PNG secara otomatis.", "Export Image")
doc.add_page_break()

# 6. GRUP: PROFIL SELESAI
# ─────────────────────────────────────────────────────────────────────────────
h_group(doc, "MODUL 6: PENGATURAN AKUN")
divider(doc)

h2(doc, "6.1 Edit Profil & Keamanan")
body(doc, "Sistem memungkinkan Anda untuk mengelola informasi pribadi dan keamanan akun melalui menu Profil.")
add_img(doc, "dropdown_profile.png", "Dropdown Profil di Header")
h3(doc, "Akses Menu Profil:")
bullet(doc, "Klik foto profil Anda di pojok kanan atas untuk memunculkan menu pilihan.", "Menu Dropdown")
filter_detail(doc, "Profile", "Masuk ke halaman pengaturan detil untuk mengubah nama, email, dan foto profil.")
filter_detail(doc, "Logout", "Keluar dari sistem secara aman. Pastikan untuk selalu logout setelah sesi selesai (tombol merah).")

add_img(doc, "15_profil.png", "Halaman Pengaturan Akun Detil")
h3(doc, "Panduan Perubahan di Halaman Profil:")
bullet(doc, "Anda bisa mengedit Nama, Email, dan Foto Profil langsung di sini.", "Update Identitas")
bullet(doc, "Sistem mendukung pengambilan foto langsung via kamera laptop/HP.", "Camera Integration")
bullet(doc, "Ganti password secara berkala untuk menjaga keamanan akun Anda.", "Update Password")
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  PENUTUP
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "PENUTUP")
divider(doc)
body(doc, "Seluruh modul di atas telah disusun berdasarkan fitur terkini pada Website Monitoring NUSTECH. Pengguna disarankan untuk selalu memeriksa data secara berkala guna memastikan validitas laporan di lapangan.")
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run(f"Mataram, {datetime.datetime.now().strftime('%d %B %Y')}\nTim Pengembang NUSTECH Monitoring")
r.font.bold = True; r.font.color.rgb = RGBColor(0x07,0x11,0x52)

# SAVE
out = r"d:\NEWNUSTECH\Panduan_Penggunaan_Website_NUSTECH_User.docx"
doc.save(out)
print(f"✅ Success: Comprehensive Manual Generated at {out}")
